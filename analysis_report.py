import psycopg2
import requests
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from db_config import DB_CONFIG

def get_stats():
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("SELECT COUNT(*) FROM apple_ukraine WHERE price_uah>0")
    total = cur.fetchone()[0]
    cur.execute("SELECT shop, COUNT(*), MIN(price_uah), MAX(price_uah), AVG(price_uah)::BIGINT FROM apple_ukraine WHERE price_uah>0 GROUP BY shop ORDER BY COUNT(*) DESC")
    shops = cur.fetchall()
    cur.execute("SELECT product, MIN(price_uah), MAX(price_uah), AVG(price_uah)::BIGINT FROM apple_ukraine WHERE price_uah>0 GROUP BY product ORDER BY AVG(price_uah) DESC")
    products = cur.fetchall()
    cur.execute("SELECT DISTINCT ON (product) product, shop, price_uah FROM apple_ukraine WHERE price_uah>0 ORDER BY product, price_uah ASC")
    cheapest = cur.fetchall()
    cur.close()
    conn.close()
    return total, shops, products, cheapest

def ask_mistral(prompt):
    r = requests.post("http://localhost:11434/api/generate", json={"model":"mistral","prompt":prompt,"stream":False}, timeout=120)
    return r.json()["response"]

def create_report():
    total, shops, products, cheapest = get_stats()
    doc = Document()

    # Заголовок
    title = doc.add_heading("Аналітичний звіт: Ринок Apple техніки в Україні", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Дата складання: 28 квітня 2026 року")
    doc.add_paragraph(f"Кількість проаналізованих позицій: {total}")
    doc.add_paragraph(f"Джерела даних: {', '.join([s[0] for s in shops])}")
    doc.add_paragraph("")

    # Розділ 1
    doc.add_heading("1. Загальна характеристика ринку", 1)
    stats_text = f"Всього зібрано {total} унікальних товарних позицій. "
    for s in shops:
        stats_text += f"Магазин {s[0]}: {s[1]} товарів, мін ціна {s[2]:,} грн, макс ціна {s[3]:,} грн, середня ціна {s[4]:,} грн. "
    prompt1 = f"Ти аналітик ринку електроніки. На основі даних напиши розгорнутий аналіз (3-4 абзаци) загальної характеристики українського ринку Apple техніки українською мовою. Дані: {stats_text}"
    print("Генерую розділ 1...")
    doc.add_paragraph(ask_mistral(prompt1))

    # Розділ 2
    doc.add_heading("2. Аналіз цін по категоріях товарів", 1)
    products_text = ""
    for p in products[:15]:
        products_text += f"{p[0]}: мін {p[1]:,} грн, макс {p[2]:,} грн, середня {p[3]:,} грн. "
    prompt2 = f"Напиши детальний аналіз (4-5 абзаців) цінових діапазонів різних категорій Apple техніки в Україні українською мовою. Порівняй iPhone, MacBook, iPad, AirPods. Дані по товарах: {products_text}"
    print("Генерую розділ 2...")
    doc.add_paragraph(ask_mistral(prompt2))

    # Розділ 3 - таблиця магазинів
    doc.add_heading("3. Порівняльна таблиця магазинів", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Магазин"
    hdr[1].text = "Кількість"
    hdr[2].text = "Мін ціна"
    hdr[3].text = "Макс ціна"
    hdr[4].text = "Середня ціна"
    for s in shops:
        row = table.add_row().cells
        row[0].text = s[0]
        row[1].text = str(s[1])
        row[2].text = f"{s[2]:,} грн"
        row[3].text = f"{s[3]:,} грн"
        row[4].text = f"{s[4]:,} грн"

    # Розділ 4 - найкращі ціни
    doc.add_heading("4. Найкращі ціни по кожній моделі", 1)
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = "Table Grid"
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "Модель"
    hdr2[1].text = "Магазин"
    hdr2[2].text = "Ціна"
    for c in cheapest:
        row = table2.add_row().cells
        row[0].text = c[0]
        row[1].text = c[1]
        row[2].text = f"{c[2]:,} грн"

    # Розділ 5
    doc.add_heading("5. Висновки та рекомендації", 1)
    prompt5 = f"На основі аналізу {total} товарів Apple в українських магазинах напиши розгорнуті висновки (4-5 абзаців) з рекомендаціями для покупців і для бізнесу українською мовою. Включи висновки про доступність техніки, конкуренцію між магазинами та перспективи ринку."
    print("Генерую розділ 5...")
    doc.add_paragraph(ask_mistral(prompt5))

    filename = "/home/vostrum/parsers/apple_market_analysis.docx"
    doc.save(filename)
    print(f"Звіт збережено: {filename}")

if __name__ == "__main__":
    create_report()
